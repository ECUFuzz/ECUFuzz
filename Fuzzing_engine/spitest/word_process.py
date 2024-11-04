from docx import Document
import re

def remove_blank_paragraphs(doc):
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            p = paragraph._element
            p.getparent().remove(p)
            p._element = p._p = None

def normalize_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def is_table_to_delete(table):
    try:
        first_cell_text = normalize_text(table.cell(0, 0).text)
        if first_cell_text != "CodeBeamer reference:":
            return False

        content = [
            ("CodeBeamer reference:", re.compile(r"SafingManagement-\d+")),
            ("Revision:", re.compile(r"\d+")),
            ("Priority:", "--"),
            ("Severity:", "--"),
            ("Status:", "Init")
        ]

        for i, (header, value) in enumerate(content):
            cell_header = normalize_text(table.cell(i, 0).text)
            cell_value = normalize_text(table.cell(i, 1).text)
            if cell_header != header:
                return False
        return True
    except IndexError:
        return False

def delete_paragraphs_with_text(doc, text):
    paragraphs_to_delete = [p for p in doc.paragraphs if text in p.text]
    for paragraph in paragraphs_to_delete:
        p = paragraph._element
        p.getparent().remove(p)
    print(f"Deleted {len(paragraphs_to_delete)} paragraphs containing '{text}'.")

def delete_tables_by_heading(doc, headings):
    for paragraph in doc.paragraphs:
        for heading in headings:
            if heading in paragraph.text:
                next_element = paragraph._element.getnext()
                if next_element is not None and next_element.tag.endswith('tbl'):
                    table_to_delete = next_element
                    table_to_delete.getparent().remove(table_to_delete)

doc_path = "64257_CHERY_T26_RCS_SC2_2S_-_DES_SW_CHY_T26_M1E_2S_SafingManagement_SRS.docx"
doc = Document(doc_path)

texts_to_delete = ["No Upstream References", "No Downstream References", "No Associations", "Description"]
for text in texts_to_delete:
    delete_paragraphs_with_text(doc, text)

headings = ["Upstream References", "Downstream References", "Associations"]
delete_tables_by_heading(doc, headings)

tables_to_delete = [table for table in doc.tables if is_table_to_delete(table)]

texts_to_delete = ["Upstream References", "Downstream References", "Associations"]
for text in texts_to_delete:
    delete_paragraphs_with_text(doc, text)

for table in tables_to_delete:
    tbl = table._element
    tbl.getparent().remove(tbl)

remove_blank_paragraphs(doc)

output_path = "modified_document.docx"
doc.save(output_path)

print(f"Deleted {len(tables_to_delete)} matching tables.")
